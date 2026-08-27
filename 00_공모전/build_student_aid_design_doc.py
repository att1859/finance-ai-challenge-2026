from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(__file__).with_name('학자금_소비평탄화_AI_서비스_설계서.docx')

NAVY = '172033'
BLUE = '2E5B88'
TEAL = '286B63'
ACID = 'C7F000'
INK = '222222'
MUTED = '667085'
LIGHT = 'F2F4F7'
PALE = 'F6F8FA'
WHITE = 'FFFFFF'
RED = '9B1C1C'
GOLD = '7A5A00'
FONT = 'Malgun Gothic'  # narrative_proposal의 한국어 가독성용 명명된 서체 override


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(total))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.first_child_found_in('w:tblInd')
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), str(indent))
    tbl_ind.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa)-1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in('w:tcW')
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(width))
            tc_w.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=None, bold=None, color=INK, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    rid = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), rid)
    run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:ascii'), FONT)
    rfonts.set(qn('w:hAnsi'), FONT)
    rfonts.set(qn('w:eastAsia'), FONT)
    rpr.append(rfonts)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rpr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rpr.append(u)
    run.append(rpr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run('Page ')
    set_run(r, size=8.5, color=MUTED)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)


def add_numbering(doc, num_id, bullet=True):
    numbering = doc.part.numbering_part.element
    abstract_id = 100 + num_id
    abstract = OxmlElement('w:abstractNum')
    abstract.set(qn('w:abstractNumId'), str(abstract_id))
    multi = OxmlElement('w:multiLevelType')
    multi.set(qn('w:val'), 'singleLevel')
    abstract.append(multi)
    lvl = OxmlElement('w:lvl')
    lvl.set(qn('w:ilvl'), '0')
    start = OxmlElement('w:start')
    start.set(qn('w:val'), '1')
    lvl.append(start)
    num_fmt = OxmlElement('w:numFmt')
    num_fmt.set(qn('w:val'), 'bullet' if bullet else 'decimal')
    lvl.append(num_fmt)
    lvl_text = OxmlElement('w:lvlText')
    lvl_text.set(qn('w:val'), '•' if bullet else '%1.')
    lvl.append(lvl_text)
    suff = OxmlElement('w:suff')
    suff.set(qn('w:val'), 'tab')
    lvl.append(suff)
    ppr = OxmlElement('w:pPr')
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'num')
    tab.set(qn('w:pos'), '540')
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '540')
    ind.set(qn('w:hanging'), '280')
    ppr.append(ind)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '80')
    spacing.set(qn('w:line'), '290')
    spacing.set(qn('w:lineRule'), 'auto')
    ppr.append(spacing)
    lvl.append(ppr)
    numbering.append(abstract)
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))
    abstract_ref = OxmlElement('w:abstractNumId')
    abstract_ref.set(qn('w:val'), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)


def add_list_item(doc, text, num_id=11):
    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    num_pr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    num = OxmlElement('w:numId')
    num.set(qn('w:val'), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    ppr.append(num_pr)
    r = p.add_run(text)
    set_run(r, size=10.5)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run(r1, size=10.5, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run(r2, size=10.5)
    else:
        r = p.add_run(text)
        set_run(r, size=10.5)
    return p


def add_callout(doc, title, body, fill='F4F6F9', accent=TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run(r, size=11, bold=True, color=accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run(r2, size=10.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    set_repeat_table_header(table.rows[0])
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], LIGHT)
        p = table.rows[0].cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run(r, size=9.2, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run(r, size=9.0)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    r = p.add_run(text)
    set_run(r, size={1:16, 2:13, 3:11.5}[level], bold=True, color=BLUE if level < 3 else TEAL)
    return p


def page_break(doc):
    # 본문 장은 자연스럽게 이어지게 하여 빈 페이지와 한 줄짜리 페이지를 막는다.
    return None


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.85)
sec.bottom_margin = Inches(0.85)
sec.left_margin = Inches(1.0)
sec.right_margin = Inches(1.0)
sec.header_distance = Inches(0.45)
sec.footer_distance = Inches(0.45)

styles = doc.styles
normal = styles['Normal']
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn('w:ascii'), FONT)
normal._element.rPr.rFonts.set(qn('w:hAnsi'), FONT)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.25

for lvl, size, before, after, color in [(1,16,16,8,BLUE),(2,13,12,6,BLUE),(3,11.5,8,4,TEAL)]:
    st = styles[f'Heading {lvl}']
    st.font.name = FONT
    st._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

add_numbering(doc, 11, bullet=True)
add_numbering(doc, 12, bullet=False)

header = sec.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = hp.add_run('2026 금융 AI Challenge | 서비스 기준 설계서')
set_run(hr, size=8.5, color=MUTED)
footer = sec.footer
add_page_field(footer.paragraphs[0])

# Cover — proposal_centerpiece
doc.add_paragraph().paragraph_format.space_after = Pt(44)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('2026 금융 AI Challenge')
set_run(r, size=11, bold=True, color=TEAL)
p.paragraph_format.space_after = Pt(16)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('학자금 소비평탄화 AI')
set_run(r, size=27, bold=True, color=NAVY)
p.paragraph_format.space_after = Pt(6)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('장학금·근로·학자금대출을 함께 설계하는 개인화 금융 시뮬레이터')
set_run(r, size=13, color=BLUE)
p.paragraph_format.space_after = Pt(26)
add_callout(doc, '핵심 제안', '상환 의무가 없는 지원을 먼저 찾고, 남은 총필요자금을 근로와 공적 학자금대출로 배치하여 대학 시절의 생활·학업시간과 졸업 후 상환 부담을 함께 평탄화한다.', fill='F1F5F4')
doc.add_paragraph().paragraph_format.space_after = Pt(24)
meta = add_table(doc, ['문서 성격', '대상', '기준일'], [['공모전 제출 준비용 기준 설계서', '국내 학부 재학생', '2026-08-27']], [3120,3120,3120])
doc.add_paragraph().paragraph_format.space_after = Pt(32)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('주의: 본 문서는 제품 설계 및 공모전 검증을 위한 자료이며, 금융·법률 자문이나 장학금·대출 승인 결과가 아니다.')
set_run(r, size=9, color=MUTED, italic=True)

doc.add_page_break()
add_heading(doc, '0. 의사결정 요약', 1)
add_callout(doc, '제품 한 문장 정의', '장학금과 감면 혜택을 먼저 찾고, 남은 필요금액 안에서 학업시간과 졸업 후 상환 부담을 함께 비교하여 현재와 미래의 총지출이 한쪽으로 무너지지 않도록 선택지를 제시한다.')
for item in [
    '첫 고객은 국내 학부 재학생이며, 비회원 상태에서도 핵심 계산을 완료할 수 있다.',
    '성과는 대출액 증가가 아니라 발견한 장학금, 근로시간 감소 가능분, 총지출 격차 감소, 상환 안전성, 추천 이해도로 측정한다.',
    '한국장학재단을 범위로 선택하면 해당 학기 국내 학부생 대상 공개 지원사업 전체를 목록에 포함한다.',
    '공모전 MVP는 수기 입력과 명확히 표시된 샘플 마이데이터 불러오기를 제공하며 실제 금융정보는 수집하지 않는다.',
    '규칙엔진이 자격 후보를 계산하고 최적화 엔진이 조합을 산출하며, AI는 공고 구조화와 설명을 담당한다.',
    '공식 자격·승인·신청은 한국장학재단과 각 기관의 최종 판단으로 남긴다.'
]: add_list_item(doc, item)

add_heading(doc, '공모전 제출과의 정합성', 2)
add_table(doc, ['공식 요구', '본 프로젝트의 대응'], [
    ['금융 현안 정의', '대학 시절의 낮은 소득·과도한 근로와 취업 후 상환 부담 사이의 불균형'],
    ['AI 기반 해결', '비정형 공고 구조화, 누락·예외 탐지, 개인별 결과 설명'],
    ['작동 가능한 웹 MVP', '샘플 프로필·대표 규칙·세 시나리오·스트레스 테스트가 실제 동작'],
    ['기획서 PDF', '본 기준 문서를 공식 양식에 맞춰 축약'],
    ['기능명세서 PDF', '본 문서의 기능표·데이터 흐름·예외처리를 공식 양식에 이관'],
    ['실행 URL', '심사 확인 기간 동안 공개 접근 가능한 정적 웹 배포']
], [2100,7260])

page_break(doc)
add_heading(doc, '1. 문제 정의와 소비 평탄화 논리', 1)
add_body(doc, '대학생은 졸업 전 소득이 낮고 시간 제약이 크지만, 취업 후에는 상대적으로 소득이 높아질 가능성이 있다. 이때 저금리 학자금대출은 단순한 비용이 아니라 현재의 결핍과 과도한 근로를 미래소득으로 일부 이전하는 금융수단이 될 수 있다. 단, 미래소득의 불확실성, 장학금 기회, 대출상품의 실제 용도와 상환조건을 함께 고려해야 한다.')
add_heading(doc, '1.1 제품이 평탄화하는 값', 2)
add_callout(doc, '용어 결정', '사용자 요청에 따라 등록금도 총소비 관점에 포함한다. 다만 화면과 계산에서는 오해를 막기 위해 이를 총필요자금으로 표시하고 교육비와 생활소비를 분해한다.', fill='FFF8E8', accent=GOLD)
add_table(doc, ['구성요소', '정의', '처리 방식'], [
    ['교육비', '등록금과 교육 관련 비용', '장학금·감면·등록금대출로 우선 충당'],
    ['생활소비', '주거·식비·교통·통신·여가 등 월 지출', '생활비지원·근로소득·생활비대출로 충당'],
    ['총필요자금', '교육비와 생활소비의 합', '대학 기간 전체 자금격차 계산에 사용'],
    ['취업 후 가처분소득', '세후소득에서 필수지출과 상환액을 차감한 값', '상환 안전성과 미래 생활수준 판단에 사용']
], [1600,3500,4260])
add_heading(doc, '1.2 최적화의 목적과 제약', 2)
add_body(doc, '목적함수는 현재와 미래의 총지출 격차, 대학 시절 노동 피로, 상환 부담, 취업·소득 불확실성을 동시에 줄이는 다목적 구조로 정의한다. 단일한 “최적 대출액”을 확정하지 않고 가치관이 다른 세 가지 해를 제공한다.')
for item in [
    '상환·근로 의무 없는 장학금과 감면은 대출보다 먼저 적용한다.',
    '근로장학금은 금액뿐 아니라 필요한 근로시간과 학업시간 손실을 함께 계산한다.',
    '상품별 등록금·생활비 용도, 한도, 금리, 상환방식을 실제 규칙대로 구분한다.',
    '장학금 미선정, 취업 6·12개월 지연, 초봉 20% 감소, 휴학·졸업 지연을 스트레스 테스트한다.',
    '필수생활비 또는 사용자가 정한 상환 안전선을 침범하는 고대출안은 추천하지 않는다.'
]: add_list_item(doc, item)

page_break(doc)
add_heading(doc, '2. 고객, 가치제안, 서비스 경계', 1)
add_heading(doc, '2.1 핵심 고객', 2)
add_body(doc, '첫 버전은 국내 학부 재학생을 대상으로 한다. 사용자는 등록금과 생활소비를 충당하기 위해 장학금, 가족지원, 근로소득, 학자금대출을 조합해야 하지만 각 수단이 학업시간과 졸업 후 생활에 미치는 영향을 한 번에 비교하기 어렵다.')
add_heading(doc, '2.2 제공 가치', 2)
add_table(doc, ['사용자의 질문', '서비스가 제공하는 답'], [
    ['받을 수 있는 지원이 무엇인가?', '한국장학재단 전체 목록과 대표 교내·지역 후보, 매칭·제외 이유'],
    ['그래도 얼마나 부족한가?', '장학금 적용 후 총필요자금과 기간별 자금격차'],
    ['얼마나 일해야 하는가?', '현재·희망 근로시간과 확보 가능한 학업시간'],
    ['대출을 얼마나 써도 되는가?', '세 가치관별 권장 범위와 총이자·상환부담'],
    ['취업이 늦어지면 어떻게 되는가?', '취업지연·초봉하락·장학금 미선정 시나리오']
], [3000,6360])
add_heading(doc, '2.3 하지 않는 일', 2)
for item in [
    '전국 모든 민간·지자체 장학금을 완전하게 제공한다고 주장하지 않는다.',
    '장학금 수혜, 학자금대출 승인 또는 신용평가를 확정하지 않는다.',
    '금융기관·한국장학재단의 로그인 정보나 공동인증서를 수집하지 않는다.',
    '자동 신청, 증빙서류 장기보관, 자기소개서 자동작성, 범용 상담 챗봇을 MVP에 포함하지 않는다.',
    '일반 신용대출을 소비 평탄화 수단으로 추천하지 않는다.'
]: add_list_item(doc, item)

page_break(doc)
add_heading(doc, '3. 서비스 여정과 화면 구조', 1)
add_heading(doc, '3.1 사용자 여정', 2)
steps = [
    ('1. 시작', '소비 평탄화를 쉬운 말로 설명하고 빠른 진단과 상세 설계를 구분한다.'),
    ('2. 금융현황', '수기 입력 또는 공모전용 샘플 마이데이터로 기존 대출·월소득·고정비·소비평균을 채운다.'),
    ('3. 학적·장학조건', '학교·학년·등록금·지원구간·성적·지역·특별자격을 필요한 순서대로 질문한다.'),
    ('4. 전체 지원사업', '한국장학재단 전체 목록에서 자동매칭·추가심사·정보부족·대상아님·모집종료를 구분한다.'),
    ('5. 개인 목표', '현재·희망 근로시간, 총필요자금, 예상 취업시점과 초봉 범위를 입력한다.'),
    ('6. 세 시나리오', '학업시간 확보형·균형형·부채 최소형을 동일한 지표로 비교한다.'),
    ('7. 위험 확인', '장학금 미선정·취업 지연·초봉 감소·졸업 지연을 즉시 재계산한다.'),
    ('8. 공식 연결', '공식 신청 페이지와 문의처로 이동하며 최종 판정은 기관에 있음을 고지한다.')
]
add_table(doc, ['단계', '핵심 동작'], steps, [1700,7660])
add_heading(doc, '3.2 결과 화면의 고정 지표', 2)
for item in [
    '장학금: 확정이 아니라 조건상 후보 금액과 판정 신뢰도',
    '총필요자금: 교육비와 생활소비의 분해',
    '시간: 주당 근로시간과 확보되는 학업시간',
    '부채: 상품별 원금, 금리, 총이자, 상환개시, 월 상환액',
    '미래: 취업 후 가처분소득과 상환 종료 후 변화',
    '근거: 사용 입력, 계산식, 기준 학기, 공식 출처, 최종 검수일'
]: add_list_item(doc, item)

page_break(doc)
add_heading(doc, '4. 장학금·학자금대출 규정 출처와 완전성', 1)
add_heading(doc, '4.1 출처 우선순위', 2)
add_table(doc, ['순위', '출처', '역할', '갱신'], [
    ['1', '국가법령정보센터', '법률·시행령·고시·시행일·개정 이력', '매일 변경 확인'],
    ['2', '한국장학재단', '국가장학·근로·기부장학·대출의 학기별 규칙', '학기 전 매일, 평시 주 1회'],
    ['3', '교육부', '연도별 기본계획과 제도 변경', '연간·수시'],
    ['4', '대학알리미·공공데이터포털', '등록금·대학정보·과거 수혜통계', '공시 주기'],
    ['5', '개별 대학·지자체·민간재단', '교내·지역·민간 공고', '모집기간 중 매일']
], [600,2200,4260,2300])
add_heading(doc, '4.2 한국장학재단 범위 계약', 2)
add_callout(doc, '완전성 원칙', '해당 학기 국내 학부생 대상 한국장학재단 공개 지원사업 전체를 목록에 포함한다. 자동 판정이 어려운 사업도 숨기지 않고 추가심사 필요 또는 정보 부족 상태로 노출한다.')
for item in [
    '사업별 program_id, 기관, 연도·학기, 규칙 버전, 시행일과 원문 해시를 저장한다.',
    '등록금성·생활비성·근로대가성·이자지원 여부를 분리한다.',
    '학적·학년·지원구간·성적·지역·특별자격·수혜횟수·중복 가능성을 구조화한다.',
    '기존 규정을 새 규정으로 덮어쓰지 않고 불변 스냅샷으로 보존한다.',
    '지원금·금리·상환기준소득·한도·신청일은 고위험 필드로 이중 검수한다.',
    '추천 결과에서 전체 지원사업 보기와 제외 이유를 제공한다.'
]: add_list_item(doc, item)
add_heading(doc, '4.3 공공데이터의 한계', 2)
add_body(doc, '공공데이터포털의 대학별 등록금·장학금 수혜 통계는 기준값과 환경 분석에는 유용하지만, 현재 신청 가능한 장학금의 자격·마감일을 판정하는 단일 통합 API가 아니다. 대학·지자체·민간 공고는 기관별 공식페이지와 제휴 피드를 사용하고, 데이터 범위와 최종 확인일을 공개한다.')

page_break(doc)
add_heading(doc, '5. 고객정보 입력, 관리, 활용', 1)
add_heading(doc, '5.1 MVP 개인정보 원칙', 2)
add_callout(doc, 'MVP 기본값', '계정 없이 브라우저에서 계산하고 입력값을 서버로 전송하지 않는다. 세션이 끝나면 삭제하며, 샘플 마이데이터는 실제 고객정보가 아닌 명시된 가상 데이터다.', fill='F1F5F4')
add_table(doc, ['구분', '항목', '처리'], [
    ['정확값', '등록금·월 생활소비·근로시간·시급·기존 학자금대출 잔액', '브라우저 세션에서만 계산'],
    ['구간값', '가족지원·소득·자산·예상초봉', '최소 단위로 입력'],
    ['선택값', '지원구간·다자녀·지역·특별자격', '필요 이유를 표시하고 선택 입력'],
    ['금지', '주민등록번호·계좌번호·인증서·금융기관 비밀번호', '수집하지 않음']
], [1200,4700,3460])
add_heading(doc, '5.2 향후 저장형 서비스', 2)
add_table(doc, ['영역', '권고 구조'], [
    ['식별정보', '인증 저장소와 계산 프로필을 가명 사용자 ID로 분리'],
    ['민감자격', '별도 저장소·별도 선택동의·최소권한 접근·학기별 재확인'],
    ['보안', 'TLS, 저장 암호화, 비밀번호 일방향 해시, 관리자 MFA, RBAC, 접속기록'],
    ['AI', '원본 개인정보 미전송, 비식별·집계 결과만 설명용 모델에 전달'],
    ['학습', '고객 입력·상담·추천 결과를 기본적으로 모델 학습에 사용하지 않음'],
    ['삭제', '탈퇴·목적달성 시 파기, 백업·로그·벡터 저장소까지 삭제 흐름 관리']
], [1900,7460])
add_heading(doc, '5.3 허용 목적과 금지 목적', 2)
add_table(doc, ['허용', '금지'], [
    ['장학금 후보·소비평탄화 계산', '대출 광고 타기팅'],
    ['선택 시 프로필·계획 저장', '신용위험 평가와 가격차별'],
    ['선택 시 신청기한·규정 변경 알림', '외부 AI 학습과 제3자 판매'],
    ['익명 집계 기반 오류·UX 개선', '취약계층 정보를 다른 금융상품에 재사용']
], [4680,4680])

page_break(doc)
add_heading(doc, '6. 기존 금융 채널과 신용정보 연동', 1)
add_heading(doc, '6.1 공모전 MVP의 표현', 2)
add_body(doc, '공식 대회는 금융소비자 특성과 서비스 채널을 고려한 웹서비스를 요구한다. MVP에는 기존 금융채널 활용 가능성을 보여주되 실제 금융정보를 수집하는 것처럼 오인시키지 않는다.')
add_table(doc, ['구분', 'MVP', '향후 상용 서비스'], [
    ['입력', '수기 입력 + 샘플 마이데이터 불러오기', '인가된 마이데이터 사업자 또는 금융회사 API 제휴'],
    ['데이터', '가상 대출잔액·월소득·고정비·소비평균', '사용자 동의 범위의 집계 금융정보'],
    ['저장', '브라우저 세션 종료 시 삭제', '선택동의·보유기간·철회·파기 관리'],
    ['배포 채널', '공개 웹 MVP', '은행 앱·마이데이터·대학 학생지원 채널 내 임베드 가능'],
    ['금지', '실제 로그인·비밀번호·계좌정보 수집 없음', '무허가 스크래핑 또는 무허가 신용정보 통합 없음']
], [1500,3730,4130])
add_heading(doc, '6.2 연동 후 사용할 최소 정보', 2)
for item in [
    '기존 학자금대출 상품·잔액·월 상환액',
    '월소득과 정기적 가족지원의 집계값',
    '최근 소비의 필수·선택 항목별 월평균',
    '기존 고정비와 다른 채무의 월 상환 합계',
    '사용자가 직접 수정하고 연동을 해제할 수 있는 권리'
]: add_list_item(doc, item)
add_body(doc, '금융회사·공공기관의 개인신용정보를 자동 수집·분석하여 제공하는 단계는 본인신용정보관리업 등 별도 규제 검토가 필요하다. 초기에는 공식 페이지로 연결하고 개인정보를 기관에 전달하지 않는 구조를 유지한다.')

page_break(doc)
add_heading(doc, '7. AI·규칙엔진·최적화엔진의 역할', 1)
add_table(doc, ['단계', '담당', '핵심 통제'], [
    ['공고 수집', 'AI 보조', '비정형 공고에서 조건·금액·기한·서류를 구조화'],
    ['규칙 승인', '사람 검수', '원문 대조·버전·라이선스·고위험 필드 이중검수'],
    ['자격 후보', '규칙엔진', '결정 규칙과 제외 이유를 재현 가능하게 계산'],
    ['시나리오 생성', '최적화엔진', '장학금·근로·대출 조합과 제약조건 계산'],
    ['설명', 'AI', '비식별 결과를 쉬운 말로 요약하고 누락정보 질문'],
    ['최종 결정', '사용자·공식기관', '사용자가 선택하고 기관이 자격·승인을 확정']
], [1600,1900,5860])
add_heading(doc, '7.1 AI가 필요한 이유', 2)
for item in [
    '장학금 공고는 기관마다 형식이 달라 규칙의 구조화와 변경점 비교에 비용이 많이 든다.',
    '사용자별 추천은 동일한 숫자라도 생활·학업·부채 우선순위에 따라 설명이 달라야 한다.',
    '정보가 부족하거나 규칙이 모호할 때 확정 판정 대신 추가 질문과 공식 확인 경로를 안내해야 한다.'
]: add_list_item(doc, item)
add_heading(doc, '7.2 AI가 결정하지 않는 것', 2)
for item in [
    '장학금 최종 수혜 여부',
    '학자금대출 승인·거절 또는 신용등급',
    '개인의 취업 가능성과 소득을 단일값으로 단정하는 예측',
    '고객 원본정보의 재사용·광고 타기팅·가격결정'
]: add_list_item(doc, item)

page_break(doc)
add_heading(doc, '8. 추천 결과와 위험 시나리오', 1)
add_heading(doc, '8.1 세 가지 결과', 2)
add_table(doc, ['시나리오', '최적화 방향', '표시'], [
    ['학업시간 확보형', '상환 안전선 안에서 근로시간을 최소화', '확보 시간과 미래 상환 부담을 함께 강조'],
    ['균형형 (기본)', '총지출 격차·근로·부채 부담의 종합 개선', '사용자 우선순위에 가장 가까운 기준안'],
    ['부채 최소형', '장학금과 감당 가능한 근로를 우선하고 대출 최소화', '대학 시절 생활수준 저하 가능성을 함께 표시']
], [1800,3800,3760])
add_heading(doc, '8.2 취업소득 입력', 2)
add_body(doc, '학교·성별·지역으로 개인 취업 가능성을 단정하지 않는다. 사용자가 예상 취업시점과 초봉 구간을 입력하고, 공식 직업·고용 통계는 참고범위로만 제공한다.')
add_table(doc, ['전망', '가정'], [
    ['보수적', '취업 12개월 지연 + 기준 초봉의 80%'],
    ['기준', '사용자가 입력한 취업시점과 초봉'],
    ['낙관적', '즉시 취업 + 기준 초봉의 120%']
], [1800,7560])
add_heading(doc, '8.3 안전 표시', 2)
for item in [
    '정확한 단일 금액 대신 권장 범위와 가정 민감도를 제시한다.',
    '고대출안을 초록색·왕관·AI 추천 배지로 강조하지 않는다.',
    '월 상환액뿐 아니라 총이자·총상환액·상환개시 조건·상환 종료 후 변화를 표시한다.',
    '자격 결과는 조건상 가능성 높음·추가 확인 필요·정보 부족·대상 아님으로 구분한다.',
    '모든 결과에 공식 신청 결과가 최종 결정임을 명시한다.'
]: add_list_item(doc, item)

page_break(doc)
add_heading(doc, '9. 공모전 MVP 기능명세 요약', 1)
features = [
    ['F01', '시작·개념 설명', '총필요자금과 소비 평탄화 설명', '필수'],
    ['F02', '샘플 마이데이터', '가상 대출·소득·고정비·소비 불러오기와 수정', '필수'],
    ['F03', '개인 조건 입력', '학적·지원구간·성적·지역·소비·근로·취업 전망', '필수'],
    ['F04', 'KOSAF 전체 목록', '대상 학부 지원사업 전체와 상태·출처·기준일', '필수'],
    ['F05', '자격 사전점검', '후보·제외 이유·부족정보·추가심사 표시', '필수'],
    ['F06', '자금격차 계산', '장학금 적용 후 교육비·생활소비 부족분 계산', '필수'],
    ['F07', '세 시나리오', '학업시간 확보·균형·부채 최소 비교', '필수'],
    ['F08', '스트레스 테스트', '장학금 미선정·취업지연·초봉감소·졸업지연', '필수'],
    ['F09', 'AI 설명', '비식별 결과 기반 추천·제외·위험 설명', '필수'],
    ['F10', '공식 연결', '원문·신청 링크·문의처·최종판정 고지', '필수'],
    ['F11', '프로필 저장', '선택 가입·동의 원장·삭제', '향후'],
    ['F12', '실제 금융연동', '마이데이터·금융회사 제휴 API', '향후']
]
add_table(doc, ['ID', '기능', '수용 기준', '범위'], features, [700,2200,5260,1200])
add_heading(doc, '9.1 대표 시연 시나리오', 2)
demo_steps = [
    '샘플 마이데이터를 불러와 기존 대출·소득·고정비·소비평균을 자동 채운다.',
    '학교·학년·등록금·지원구간·현재 근로시간을 추가한다.',
    '한국장학재단 전체 목록에서 후보와 제외 이유를 확인한다.',
    '장학금 적용 전후의 자금격차와 세 가지 소비평탄화 시나리오를 비교한다.',
    '취업 12개월 지연 또는 장학금 미선정을 적용해 결과 변화를 확인한다.',
    '공식 출처와 신청 페이지로 이동한다.'
]
for item in demo_steps: add_list_item(doc, item, num_id=12)

page_break(doc)
add_heading(doc, '10. 금융소비자 보호, 개인정보, 규제 검토', 1)
add_heading(doc, '10.1 금지 표현', 2)
for item in [
    '대출받을수록 이득입니다.',
    '지금 안 빌리면 손해입니다.',
    '안전한 빚 또는 무조건 갚을 수 있습니다.',
    '장학금 수혜 확정 또는 대출 승인 보장.',
    'AI가 당신의 최적 대출액을 결정했습니다.'
]: add_list_item(doc, item)
add_heading(doc, '10.2 출시 전 검토 게이트', 2)
add_table(doc, ['게이트', '확인사항'], [
    ['금융규제', '개인화 대출 비교·추천·신청 연결이 금융상품 중개·자문에 해당하는지'],
    ['신용정보', '개인신용정보 자동연계가 마이데이터 허가 또는 제휴 범위인지'],
    ['개인정보', '처리근거·최소수집·민감정보·위탁·제3자 제공·국외이전·보유기간'],
    ['자동화 결정', '결과 설명·입력 수정·이의제기·인적 검토가 필요한 수준인지'],
    ['AI', '고영향 AI 해당성, AI 사용 고지, 입출력 통제와 모델 공급자 계약'],
    ['데이터 권리', '공공누리 유형·기관 이용약관·원문 복제·크롤링·출처표시']
], [2000,7360])
add_callout(doc, '법적 성격', '본 문서의 개인정보·금융규제 내용은 제품 설계를 위한 검토안이며 법률 자문이 아니다. 실제 출시 전 개인정보보호책임자와 금융·개인정보 분야 전문가의 검토가 필요하다.', fill='FFF1F1', accent=RED)

page_break(doc)
add_heading(doc, '11. 운영지표와 단계별 로드맵', 1)
add_heading(doc, '11.1 핵심 성과지표', 2)
add_table(doc, ['지표', '의미'], [
    ['발견한 장학금 후보·금액', '대출 전에 상환 없는 지원을 얼마나 발견했는가'],
    ['근로시간 감소 가능분', '생활 안정성을 유지하면서 확보한 학업시간'],
    ['총지출 격차 감소', '대학 시절과 취업 후의 불균형 완화 정도'],
    ['상환 안전선 준수율', '추천안이 필수생활비와 사용자의 안전선을 지켰는가'],
    ['추천 이해도', '사용자가 제안된 금액과 이유를 30초 안에 설명할 수 있는가'],
    ['규칙 최신성', '고위험 필드와 전체 목록의 기준일·검수상태']
], [2600,6760])
add_heading(doc, '11.2 로드맵', 2)
add_table(doc, ['단계', '목표', '핵심 산출물'], [
    ['공모전 MVP', '문제·AI·웹서비스 가능성 검증', '샘플 마이데이터, 대표 규칙, 전체 목록 구조, 세 시나리오'],
    ['파일럿', '특정 대학·지역에서 정확도 검증', '기관별 데이터 피드, 저장형 프로필, 변경 알림'],
    ['제휴 서비스', '기존 금융·대학 채널 내 배포', '인가된 마이데이터/금융회사 API, 대학 학생지원 연계'],
    ['확장', '전국 장학금 범위 확대', '대학·지자체·민간재단 제휴와 품질 SLA']
], [1600,3200,4560])
add_heading(doc, '11.3 공모전 제작 우선순위', 2)
add_table(doc, ['비중', '작업'], [
    ['35%', '장학금·근로·대출 통합 최적화와 추천 근거'],
    ['25%', '한국장학재단 전체 목록과 규칙 버전 관리'],
    ['20%', '수기 입력과 샘플 금융정보 연동 체험'],
    ['10%', '개인정보·출처·기준일·불확실성 표시'],
    ['10%', 'UI 완성도·배포 안정성·제출 문서']
], [1400,7960])

page_break(doc)
add_heading(doc, '부록 A. 공식 출처', 1)
sources = [
    ('S1', '2026 금융 AI Challenge 공식 페이지', 'https://daker.ai/public/hackathons/2026-finance-ai-challenge'),
    ('S2', '국가법령정보 공동활용 API 안내', 'https://open.law.go.kr/LSO/openApi/openApiManual.do'),
    ('S3', '한국장학재단 취업후상환 학자금대출', 'https://www.kosaf.go.kr/ko/tuition.do?pg=tuition04_01_01'),
    ('S4', '한국장학재단 일반상환 학자금대출', 'https://www.kosaf.go.kr/ko/tuition.do?naviParam=HD&pg=tuition04_02_01'),
    ('S5', '한국장학재단 중복지원 안내', 'https://www.kosaf.go.kr/ko/notice.do?mode=view&seqNo=7771'),
    ('S6', '공공데이터포털 대학별 장학금 수혜 현황', 'https://www.data.go.kr/data/15038576/fileData.do'),
    ('S7', '공공데이터포털 대학알리미 재정 API', 'https://www.data.go.kr/data/15038392/openapi.do'),
    ('S8', '전국 대학 평균등록금 표준데이터', 'https://www.data.go.kr/data/15107738/standard.do'),
    ('S9', '개인정보보호법 제15조·제16조', 'https://www.law.go.kr/LSW/lsLinkCommonInfo.do?lsJoLnkSeq=1006183947'),
    ('S10', '개인정보보호법 제23조 민감정보', 'https://www.law.go.kr/LSW/lsLinkCommonInfo.do?lsJoLnkSeq=1027416043'),
    ('S11', '개인정보보호법 제24조의2 주민등록번호', 'https://www.law.go.kr/lsLinkCommonInfo.do?lsJoLnkSeq=1006184231'),
    ('S12', '개인정보보호법 제26조 처리위탁', 'https://www.law.go.kr/LSW/lsLawLinkInfo.do?ancYnChk=0&chrClsCd=010202&lsJoLnkSeq=1005993251'),
    ('S13', '개인정보보호법 제28조의8 국외이전', 'https://www.law.go.kr/lsLinkCommonInfo.do?chrClsCd=010202&lsJoLnkSeq=1029334953'),
    ('S14', '개인정보보호법 제37조의2 자동화된 결정', 'https://law.go.kr/LSW/lsLinkCommonInfo.do?chrClsCd=010202&lsJoLnkSeq=1029334889'),
    ('S15', '개인정보위 생성형 AI 개인정보 안내서', 'https://pipc.go.kr/np/cop/bbs/selectBoardArticle.do?bbsId=BS074&mCode=C020010000&nttId=11410'),
    ('S16', '금융위원회 마이데이터 허가 안내', 'https://www.fsc.go.kr/po010102/74324'),
    ('S17', '금융소비자보호법', 'https://www.law.go.kr/lsLinkCommonInfo.do?lsJoLnkSeq=1010081863'),
    ('S18', '공공누리 이용조건 유형', 'https://www.kogl.or.kr/etc/allMenu.do')
]
for sid, name, url in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f'[{sid}] {name}: ')
    set_run(r, size=9.2, bold=True, color=NAVY)
    add_hyperlink(p, url, url)

add_heading(doc, '부록 B. 합의된 제품 원칙', 1)
for item in [
    '공개 약속은 대출 적극 활용이 아니라 장학금·근로·대출을 조합한 생활·상환 설계다.',
    '첫 고객은 국내 학부 재학생이다.',
    '총소비에는 등록금을 포함하되 화면에서는 교육비와 생활소비를 분해한 총필요자금으로 표시한다.',
    '한국장학재단을 범위로 선택하면 해당 학기 대상 공개사업을 빠짐없이 목록화한다.',
    'MVP는 비회원 브라우저 계산이며 실제 고객 금융정보를 저장하지 않는다.',
    '기존 금융채널 연동은 샘플 마이데이터로 시연하고 실제 서비스는 인가 사업자와 제휴한다.',
    'AI는 공고 구조화와 설명을 담당하며 공식 자격·대출승인을 결정하지 않는다.',
    '기본 시나리오는 균형형이고 학업시간 확보형·부채 최소형을 함께 비교한다.'
]: add_list_item(doc, item)

doc.core_properties.title = '학자금 소비평탄화 AI 서비스 설계서'
doc.core_properties.subject = '2026 금융 AI Challenge 공모전 제출 준비용 기준 문서'
doc.core_properties.author = '학자금 소비평탄화 AI 프로젝트팀'
doc.core_properties.keywords = '학자금대출, 장학금, 소비평탄화, AI, 개인정보, 마이데이터'

doc.save(OUT)
print(OUT)
